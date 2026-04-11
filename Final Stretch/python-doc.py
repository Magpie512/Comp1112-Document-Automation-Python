from docx import Document
doc = Document()
doc.add_heading("Title", level=1)
p = doc.add_paragraph("Hello ")
p.add_run("world").bold = True
doc.save("out.docx")