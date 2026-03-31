"""
Mars Briggs 200561234
2026-03-31 11:23am
Read invoice document and then transcribe to docx file
"""
import openpyxl
import docx
import time

wb = openpyxl.load_workbook(r"C:\Lesson 11\InvoiceSheet.xlsx")
sheet = wb.active
invoiceList = []

# Columns in use
colList = ["A","B","C","D","E","F","G","H","I","J","K","L","M"]
# Jesus christ dude

# Adding data row by row
for row in range(1, 13):
    rowData = []
    for col in range(1, 14):  # 13 columns
        cellValue = sheet.cell(row=row, column=col).value
        rowData.append(cellValue)
    invoiceList.append(rowData)

print(invoiceList)

# Creates the document
doc = docx.Document()
doc.add_heading('Invoice',0)
doc.add_picture('cat.jpg', width=docx.shared.Inches(1), height=docx.shared.Cm(3))

doc.add_paragraph('Your Company Name\t\t\tINVOICE TO:\t\t\tInvoice Number')

# Company Grab
companyName = invoiceList[1][0]
companyInfo = invoiceList[1][1].split(",")
companyAddress = companyInfo[0]
companyCityState = companyInfo[1]
companyCounty = companyInfo[2]
companyPostal = companyInfo[3]

invoiceToAddr = invoiceList[1][1]
invoiceID = invoiceList[1][2]
dueDate = invoiceList[1][3]
description = invoiceList[1][4]
QTY = invoiceList[1][5]
unitPrice = invoiceList[1][6]
total = invoiceList[1][7]
taxRate = invoiceList[1][8]
taxTotal= invoiceList[1][9]
BalanceDue = invoiceList[1][10]
Notes = invoiceList[1][11]
Toss = invoiceList[1][12]

todayDate = time.time()
#dateInvoiced = (str.todayDate.year)
#print (dateInvoiced)

dueDatePrint = str(dueDate.year) +"-"+ str(dueDate.month)+"-"+ str(dueDate.day)

doc.add_paragraph(str(companyName) +"\t\t\t\t" + "Apple" + "\t\t\t\t\t" + str(invoiceID))
doc.add_paragraph(str(companyAddress) + "\t\t\t" + "123 Drive Drive" + "\t\t\t\t" + str(invoiceID))
doc.add_paragraph("\t\t\t\t\t\t\t\t\t Date of Invoice")
doc.add_paragraph("\t\t\t\t\t\t\t   " + str(time.ctime(todayDate)))
doc.add_paragraph("\t\t\t\t\t\t\t\t\t Due Date")
doc.add_paragraph("\t\t\t\t\t\t\t\t\t   " + str(dueDatePrint))
doc.add_heading('Description\t\t\t\tQTY\t\tUnit Price \t\t Total',3)
doc.save('Invoice Sheet.docx')
