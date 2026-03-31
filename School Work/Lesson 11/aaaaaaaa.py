import docx
doc = docx.Document()
doc.add_heading('Cat Photo',0)

paraObj1 = doc.add_paragraph('This is cat')

paraObj1.add_run(',yo!')

doc.add_picture('cat.jpg', width=docx.shared.Inches(2), height=docx.shared.Cm(7))
doc.save("new.docx")
